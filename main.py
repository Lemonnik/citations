import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
import working_w_doc
import pubmed_parsing
import os
import zipfile
import tempfile
import shutil
import io


some_file = "results_copied.docx"  # file we want to work with
zipfile_name = "result.zip"
result_file = "results_result.docx"  # file we want to get in the en

pmid_list = working_w_doc.main(some_file)  # getting PMIDs from WORD file WITH REPEATS(!)
pmid_list = list(dict.fromkeys(pmid_list))  # DELETE REPEATED PMIDs

print("Found " + str(len(pmid_list)) + " PMIDs!")
print("Creating links...\n")

citations = pubmed_parsing.main(pmid_list)  # getting all citations from the Internet
# citations.sort()  # if we need them in alphabetical order

doc = docx.Document(some_file)  # open file to write in
i = 0

print("\nAdding citations in the end of the file...\n[", end="")
for item in citations:
    i += 1
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(str(i) + ". " + item)
    font = run.font
    font.name = 'Times New Roman'

    font.size = Pt(14)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    print(".", end="")

print("]\nCitations added!")
doc.save(result_file)  # saving results

# now we've got file with citations in the end, but we want to replace (pmid) to [1]


# Now we have to convert .docx to .zip
# Then extract .zip archive and get word/document.xml file
# in which we will replace text
os.rename(result_file, zipfile_name)


# another black box to DELETE file in .zip
def remove_from_zip(zip_fname, *filenames):
    tempdir = tempfile.mkdtemp()
    try:
        temp_name = os.path.join(tempdir, 'new.zip')
        with zipfile.ZipFile(zip_fname, 'r') as zip_read:
            with zipfile.ZipFile(temp_name, 'w') as zip_write:
                for item in zip_read.infolist():
                    if item.filename not in filenames:
                        data = zip_read.read(item.filename)
                        zip_write.writestr(item, data)
        shutil.move(temp_name, zip_fname)
    finally:
        shutil.rmtree(tempdir)


with zipfile.ZipFile(zipfile_name, 'r') as z:
    # printing all the contents of the zip file
    print("\n\n.zip content:")
    z.printdir()
    # extracting one file
    print('\n\nExtracting word/document.xml file now...')
    z.extract('word/document.xml')
    print('Done!')
    z.close()

# read document.xml
# fin = open('word/document.xml', "rt")
# lines = fin.readlines()
# fin.close()
# READ IT in UTF-8 WAY
with io.open('word/document.xml', encoding='utf-8') as file:
    lines = file.readlines()
    file.close()

print("\nReplacing...")
for i, elem in enumerate(pmid_list):
    print("[" + str(i + 1) + "] instead of (" + str(elem) + ")")
    lines[1] = lines[1].replace('('+str(elem)+')', '[' + str(i+1) + ']')
    # print(lines[1])

print("Replacing completed!\n\nWriting to file...")

# we will write back in this file

with io.open('word/document.xml', encoding='utf-8', mode='w') as file:
    for line in lines:
        file.write(line)
        # print(line)
    file.close()


print("Writing done successfully!\n\nPacking word/document.xml back...")

# delete old document.xml and add new one to .zip
remove_from_zip(zipfile_name, "word/document.xml")
with zipfile.ZipFile(zipfile_name, "a") as z:
    # writing file
    z.write("word/document.xml")
    z.close()

print("document.xml packed successfully!")

os.rename(zipfile_name, result_file)



